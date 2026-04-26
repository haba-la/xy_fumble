import io
import json
import re
from typing import Callable, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

try:
    import json5
except Exception:
    json5 = None


ALIGNMENT_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

FONT_SIZE_MAP = {
    "初号": Pt(42),
    "小初": Pt(36),
    "一号": Pt(26),
    "小一": Pt(24),
    "二号": Pt(22),
    "小二": Pt(18),
    "三号": Pt(16),
    "小三": Pt(15),
    "四号": Pt(14),
    "小四": Pt(12),
    "五号": Pt(10.5),
    "小五": Pt(9),
    "六号": Pt(7.5),
}

MAX_PROMPT_PARAGRAPH_CHARS = 5000


def _extract_json(text: str) -> str:
    text = text.strip()
    code_fence = re.search(r"```(?:json)?\s*([\s\S]*?)```", text, re.IGNORECASE)
    if code_fence:
        return code_fence.group(1).strip()

    brace_count = 0
    start = -1
    for i, ch in enumerate(text):
        if ch == "{":
            if brace_count == 0:
                start = i
            brace_count += 1
        elif ch == "}":
            brace_count -= 1
            if brace_count == 0 and start != -1:
                return text[start : i + 1]

    return text


def _fix_trailing_comma(text: str) -> str:
    # 删除对象或数组中的尾随逗号
    text = re.sub(r",\s*([}\]])", r"\1", text)
    return text


def _parse_llm_json(raw_text: str):
    raw_json = _extract_json(raw_text)
    candidates = [raw_json]

    fixed = _fix_trailing_comma(raw_json)
    if fixed != raw_json:
        candidates.append(fixed)

    last_error = None
    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception as exc:
            last_error = exc
            if json5 is not None:
                try:
                    return json5.loads(candidate)
                except Exception as exc2:
                    last_error = exc2

    raise ValueError(f"JSON解析失败: {last_error}")


def _build_format_prompt(paragraphs: List[str], rules: Dict) -> str:
    doc_text = "\n\n".join(paragraphs)
    rules_text = json.dumps(rules, ensure_ascii=False, indent=2)
    return f"""
你是一个专业的文档排版助手。请根据文档内容语义识别每段类型，并严格按给定规则返回排版指令。

文档内容：
<doc>
{doc_text}
</doc>

排版规则：
<rules>
{rules_text}
</rules>

要求：
1. 仅返回 JSON。
2. 输出格式如下：
{{
  "elements": [
    {{
      "type": "标题",
      "content": "...",
      "format": {{
        "font": "黑体",
        "size": "小二",
        "bold": true,
        "line_spacing": 1.5,
        "alignment": "center"
      }}
    }}
  ]
}}
3. elements 数量应与输入段落数量一致，content 与原段落保持一致。
4. alignment 仅允许 left/center/right/justify。
""".strip()


def _normalize_alignment(value: str) -> str:
    if not value:
        return "left"
    value = str(value).strip().lower()
    if value in ("居中", "center", "居中对齐"):
        return "center"
    if value in ("右对齐", "right"):
        return "right"
    if value in ("两端对齐", "两端", "justify", "justified"):
        return "justify"
    return "left"


def _normalize_instruction_element(element: Dict, default_rule: Dict) -> Dict:
    fmt = dict(default_rule)
    fmt.update(element.get("format") or {})

    return {
        "type": str(element.get("type", "正文")),
        "content": str(element.get("content", "")),
        "format": {
            "font": str(fmt.get("font", "宋体")),
            "size": str(fmt.get("size", "小四")),
            "bold": bool(fmt.get("bold", False)),
            "line_spacing": float(fmt.get("line_spacing", 1.5)),
            "alignment": _normalize_alignment(fmt.get("alignment", "left")),
        },
    }


def generate_formatting_instructions(
    paragraphs: List[str],
    rules: Dict,
    llm_call: Callable[[str, str], str],
) -> Tuple[bool, Dict]:
    if not isinstance(paragraphs, list) or not paragraphs:
        return True, {"elements": []}

    system_prompt = "你是一个严谨的文档排版规划器，只输出JSON。"
    default_rule = rules.get("正文", {}) if isinstance(rules, dict) else {}

    chunks = []
    current = []
    current_chars = 0
    for p in paragraphs:
        p_len = max(len(p), 1)
        if current and current_chars + p_len > MAX_PROMPT_PARAGRAPH_CHARS:
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(p)
        current_chars += p_len
    if current:
        chunks.append(current)

    all_elements = []
    for idx, chunk in enumerate(chunks, start=1):
        prompt = _build_format_prompt(chunk, rules)
        try:
            response_text = llm_call(prompt, system_prompt=system_prompt)
        except Exception as exc:
            return False, {"error": f"调用大模型失败(分段{idx}/{len(chunks)}): {exc}"}

        try:
            parsed = _parse_llm_json(response_text)
        except Exception as exc:
            return False, {"error": f"解析排版指令失败(分段{idx}/{len(chunks)}): {exc}", "raw": response_text}

        elements = parsed.get("elements", []) if isinstance(parsed, dict) else []
        if not isinstance(elements, list) or not elements:
            return False, {"error": f"排版指令缺少 elements 列表(分段{idx}/{len(chunks)})", "raw": parsed}

        normalized_elements = [
            _normalize_instruction_element(ele if isinstance(ele, dict) else {}, default_rule)
            for ele in elements
        ]

        # 分段保底对齐
        if len(normalized_elements) < len(chunk):
            for p in chunk[len(normalized_elements):]:
                normalized_elements.append(
                    _normalize_instruction_element({"type": "正文", "content": p, "format": default_rule}, default_rule)
                )
        elif len(normalized_elements) > len(chunk):
            normalized_elements = normalized_elements[: len(chunk)]

        for i, p in enumerate(chunk):
            if not normalized_elements[i].get("content"):
                normalized_elements[i]["content"] = p

        all_elements.extend(normalized_elements)

    return True, {"elements": all_elements}


def _apply_run_format(run, fmt: Dict) -> None:
    font_name = fmt.get("font", "宋体")
    run.font.name = font_name
    size_name = fmt.get("size", "小四")
    if size_name in FONT_SIZE_MAP:
        run.font.size = FONT_SIZE_MAP[size_name]
    run.bold = bool(fmt.get("bold", False))



def _apply_paragraph_format(paragraph, fmt: Dict) -> None:
    alignment_key = _normalize_alignment(fmt.get("alignment", "left"))
    paragraph.alignment = ALIGNMENT_MAP.get(alignment_key, WD_PARAGRAPH_ALIGNMENT.LEFT)

    line_spacing = fmt.get("line_spacing", 1.5)
    try:
        paragraph.paragraph_format.line_spacing = float(line_spacing)
    except (TypeError, ValueError):
        paragraph.paragraph_format.line_spacing = 1.5


def format_docx_bytes(source_bytes: bytes, instructions: Dict) -> Tuple[bool, bytes or Dict]:
    try:
        input_doc = Document(io.BytesIO(source_bytes))
    except Exception as exc:
        return False, {"error": f"读取docx失败: {exc}"}

    elements = instructions.get("elements", []) if isinstance(instructions, dict) else []
    if not isinstance(elements, list):
        return False, {"error": "排版指令格式错误：elements 不是列表"}

    new_doc = Document()

    if not elements:
        # 保底：原样写入
        for p in input_doc.paragraphs:
            new_doc.add_paragraph(p.text)
    else:
        for element in elements:
            fmt = (element or {}).get("format", {}) if isinstance(element, dict) else {}
            content = (element or {}).get("content", "") if isinstance(element, dict) else ""

            paragraph = new_doc.add_paragraph()
            run = paragraph.add_run(str(content))
            _apply_run_format(run, fmt)
            _apply_paragraph_format(paragraph, fmt)

    out = io.BytesIO()
    new_doc.save(out)
    out.seek(0)
    return True, out.read()
