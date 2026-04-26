from __future__ import annotations

import argparse
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


BLOCK_TAGS = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "li"}
HEADING_STYLE_CANDIDATES = {
    "h1": ("Heading 1", "标题 1", "标题1", "Heading1"),
    "h2": ("Heading 2", "标题 2", "标题2", "Heading2"),
    "h3": ("Heading 3", "标题 3", "标题3", "Heading3"),
    "h4": ("Heading 4", "标题 4", "标题4", "Heading4"),
    "h5": ("Heading 5", "标题 5", "标题5", "Heading5"),
    "h6": ("Heading 6", "标题 6", "标题6", "Heading6"),
}
HEADING_TAG_LEVEL = {f"h{i}": i for i in range(1, 7)}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert HTML to DOCX while preserving layout")
    parser.add_argument("--html", type=Path, required=True, help="Input HTML file path")
    parser.add_argument("--output", type=Path, required=True, help="Output DOCX file path")
    parser.add_argument(
        "--template-docx",
        type=Path,
        default=None,
        help="Optional template DOCX path for heading style mapping",
    )
    return parser.parse_args()


def _strip_css_comments(css_text: str) -> str:
    return re.sub(r"/\*.*?\*/", "", css_text, flags=re.S)


def _parse_declarations(block: str) -> Dict[str, str]:
    result: Dict[str, str] = {}
    for part in block.split(";"):
        if ":" not in part:
            continue
        key, value = part.split(":", 1)
        k = key.strip().lower()
        v = value.strip()
        if not k or not v:
            continue
        v = re.sub(r"\s*!important\s*$", "", v, flags=re.I)
        result[k] = v
    return result


def _extract_css_rules(css_text: str) -> List[Tuple[str, Dict[str, str]]]:
    css_text = _strip_css_comments(css_text)
    rules: List[Tuple[str, Dict[str, str]]] = []
    for match in re.finditer(r"([^{}]+)\{([^{}]*)\}", css_text, flags=re.S):
        selector = match.group(1).strip()
        if selector.startswith("@"):
            continue
        decls = _parse_declarations(match.group(2))
        for one in selector.split(","):
            sel = one.strip()
            if sel:
                rules.append((sel, decls))
    return rules


def _parse_length_to_pt(value: str, base_font_pt: float = 12.0) -> Optional[float]:
    if not value:
        return None
    v = value.strip().lower()
    if v.startswith("var("):
        return None

    m = re.match(r"^(-?\d+(?:\.\d+)?)(pt|px|mm|cm|in|em|rem)?$", v)
    if not m:
        return None

    num = float(m.group(1))
    unit = m.group(2) or "pt"

    if unit == "pt":
        return num
    if unit == "px":
        return num * 72.0 / 96.0
    if unit == "mm":
        return num * 72.0 / 25.4
    if unit == "cm":
        return num * 72.0 / 2.54
    if unit == "in":
        return num * 72.0
    if unit in {"em", "rem"}:
        return num * base_font_pt
    return None


def _resolve_var(value: str, css_vars: Dict[str, str]) -> str:
    val = value.strip()
    m = re.match(r"var\((--[a-z0-9\-]+)\)", val, flags=re.I)
    if not m:
        return val
    return css_vars.get(m.group(1), val)


def _expand_4_values(values: List[str]) -> Tuple[str, str, str, str]:
    if len(values) == 1:
        return values[0], values[0], values[0], values[0]
    if len(values) == 2:
        return values[0], values[1], values[0], values[1]
    if len(values) == 3:
        return values[0], values[1], values[2], values[1]
    return values[0], values[1], values[2], values[3]


def _extract_page_and_style(html: str) -> Tuple[Dict[str, str], Dict[str, Dict[str, str]]]:
    soup = BeautifulSoup(html, "html.parser")
    css_text = "\n".join(tag.get_text("\n") for tag in soup.find_all("style"))
    css_text = _strip_css_comments(css_text)

    css_vars: Dict[str, str] = {}
    root_match = re.search(r":root\s*\{(.*?)\}", css_text, flags=re.S)
    if root_match:
        css_vars = {
            k: v
            for k, v in _parse_declarations(root_match.group(1)).items()
            if k.startswith("--")
        }

    page: Dict[str, str] = {}
    page_match = re.search(r"@page\s*\{(.*?)\}", css_text, flags=re.S)
    if page_match:
        page_decls = _parse_declarations(page_match.group(1))
        if "size" in page_decls:
            page["size"] = page_decls["size"]
        if "margin" in page_decls:
            page["margin"] = page_decls["margin"]

    style_map: Dict[str, Dict[str, str]] = {}
    for selector, decls in _extract_css_rules(css_text):
        style_map[selector] = {**style_map.get(selector, {}), **decls}

    if css_vars:
        if "size" in page:
            size_items = [_resolve_var(x, css_vars) for x in page["size"].split()]
            page["size"] = " ".join(size_items)
        if "margin" in page:
            margin_items = [_resolve_var(x, css_vars) for x in page["margin"].split()]
            page["margin"] = " ".join(margin_items)

    return page, style_map


def _selector_for_tag(tag_name: str) -> Iterable[str]:
    yield tag_name
    yield f"body {tag_name}"
    yield f"body.docx-layout-sync {tag_name}"
    yield f"body.docx-layout-sync .page {tag_name}"


def _combined_style(tag_name: str, style_map: Dict[str, Dict[str, str]]) -> Dict[str, str]:
    merged: Dict[str, str] = {}
    for selector in _selector_for_tag(tag_name):
        if selector in style_map:
            merged.update(style_map[selector])
    return merged


def _apply_page_layout(doc: Document, page: Dict[str, str]) -> None:
    section = doc.sections[0]

    size_raw = page.get("size")
    if size_raw:
        size_values = size_raw.split()
        if len(size_values) >= 2:
            width_pt = _parse_length_to_pt(size_values[0])
            height_pt = _parse_length_to_pt(size_values[1])
            if width_pt and height_pt:
                section.page_width = Mm(width_pt * 25.4 / 72.0)
                section.page_height = Mm(height_pt * 25.4 / 72.0)

    margin_raw = page.get("margin")
    if margin_raw:
        margin_values = margin_raw.split()
        if margin_values:
            t, r, b, l = _expand_4_values(margin_values)
            top = _parse_length_to_pt(t)
            right = _parse_length_to_pt(r)
            bottom = _parse_length_to_pt(b)
            left = _parse_length_to_pt(l)
            if top is not None:
                section.top_margin = Pt(top)
            if right is not None:
                section.right_margin = Pt(right)
            if bottom is not None:
                section.bottom_margin = Pt(bottom)
            if left is not None:
                section.left_margin = Pt(left)


def _set_modern_compatibility_mode(doc: Document, mode: str = "16") -> None:
    settings_el = doc.settings.element
    compat_el = settings_el.find(qn("w:compat"))
    if compat_el is None:
        compat_el = OxmlElement("w:compat")
        settings_el.append(compat_el)

    target = None
    for node in compat_el.findall(qn("w:compatSetting")):
        if node.get(qn("w:name")) == "compatibilityMode":
            target = node
            break

    if target is None:
        target = OxmlElement("w:compatSetting")
        compat_el.append(target)

    target.set(qn("w:name"), "compatibilityMode")
    target.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    target.set(qn("w:val"), mode)


def _create_output_document(template_docx_path: Optional[Path]) -> Document:
    """Create output document.

    If template DOCX is provided, use it as style base and clear body content while
    keeping style/theme definitions so heading colors and other style details persist.
    """
    if template_docx_path is None or not template_docx_path.exists():
        return Document()

    doc = Document(str(template_docx_path))
    body = doc._element.body
    # Preserve section properties, remove existing paragraphs/tables/content.
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    return doc


def _extract_style_outline_level(style) -> Optional[int]:
    ppr = style.element.find(qn("w:pPr")) if style is not None else None
    if ppr is None:
        return None
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    raw = outline.get(qn("w:val"))
    if raw is None:
        return None
    try:
        value = int(raw)
    except (TypeError, ValueError):
        return None
    return value + 1


def _normalize_hex_color(value: str) -> Optional[str]:
    if not value:
        return None
    token = value.strip().lstrip("#")
    if len(token) != 6 or not re.fullmatch(r"[0-9a-fA-F]{6}", token):
        return None
    return f"#{token.upper()}"


def _normalize_theme_color_key(value: str) -> str:
    lowered = (value or "").strip().lower().replace("_", "")
    replacements = {
        "text1": "dk1",
        "background1": "lt1",
        "text2": "dk2",
        "background2": "lt2",
        "followedhyperlink": "folhlink",
    }
    return replacements.get(lowered, lowered)


def _load_theme_color_map(docx_path: Optional[Path]) -> Dict[str, RGBColor]:
    result: Dict[str, RGBColor] = {}
    if docx_path is None or not docx_path.exists():
        return result

    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            theme_name = None
            for name in zf.namelist():
                if name.lower().startswith("word/theme/") and name.lower().endswith(".xml"):
                    theme_name = name
                    break
            if theme_name is None:
                return result

            xml_bytes = zf.read(theme_name)
            root = ET.fromstring(xml_bytes)
            ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            clr_scheme = root.find(".//a:clrScheme", ns)
            if clr_scheme is None:
                return result

            for child in list(clr_scheme):
                key = _normalize_theme_color_key(child.tag.split("}")[-1])
                hex_color = None
                srgb = child.find(".//a:srgbClr", ns)
                if srgb is not None:
                    hex_color = _normalize_hex_color(srgb.attrib.get("val", ""))
                if hex_color is None:
                    sys_clr = child.find(".//a:sysClr", ns)
                    if sys_clr is not None:
                        hex_color = _normalize_hex_color(sys_clr.attrib.get("lastClr", ""))
                if hex_color:
                    result[key] = RGBColor(
                        int(hex_color[1:3], 16),
                        int(hex_color[3:5], 16),
                        int(hex_color[5:7], 16),
                    )
    except Exception:
        return result

    return result


def _extract_rgb_from_style(style, theme_map: Dict[str, RGBColor]) -> Optional[RGBColor]:
    font = getattr(style, "font", None)
    if font is not None and getattr(font, "color", None) is not None:
        rgb = getattr(font.color, "rgb", None)
        if rgb is not None:
            return RGBColor(rgb[0], rgb[1], rgb[2])

    rpr = getattr(getattr(style, "_element", None), "rPr", None)
    if rpr is None:
        return None
    color = getattr(rpr, "color", None)
    if color is None:
        return None

    val = getattr(color, "val", None)
    if val is not None:
        normalized = _normalize_hex_color(str(val))
        if normalized:
            return RGBColor(
                int(normalized[1:3], 16),
                int(normalized[3:5], 16),
                int(normalized[5:7], 16),
            )

    theme_color = color.get(qn("w:themeColor")) if hasattr(color, "get") else None
    if theme_color:
        return theme_map.get(_normalize_theme_color_key(str(theme_color)))
    return None


def _build_heading_style_candidates_from_template(template_docx_path: Optional[Path]) -> Dict[str, Tuple[str, ...]]:
    mapping = {tag: list(names) for tag, names in HEADING_STYLE_CANDIDATES.items()}
    if template_docx_path is None or not template_docx_path.exists():
        return {tag: tuple(names) for tag, names in mapping.items()}

    template_doc = Document(str(template_docx_path))
    for style in template_doc.styles:
        if getattr(style, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
            continue

        style_name = getattr(style, "name", "")
        if not style_name:
            continue

        level = _extract_style_outline_level(style)
        if level is None:
            lowered = style_name.strip().lower()
            m = re.search(r"heading\s*([1-6])", lowered)
            if m:
                level = int(m.group(1))
            else:
                m_cn = re.search(r"标题\s*([1-6])", style_name)
                if m_cn:
                    level = int(m_cn.group(1))

        if level is None or level < 1 or level > 6:
            continue

        tag = f"h{level}"
        if style_name not in mapping[tag]:
            mapping[tag].insert(0, style_name)

    return {tag: tuple(names) for tag, names in mapping.items()}


def _extract_heading_color_map_from_template(template_docx_path: Optional[Path]) -> Dict[int, RGBColor]:
    color_map: Dict[int, RGBColor] = {}
    if template_docx_path is None or not template_docx_path.exists():
        return color_map

    template_doc = Document(str(template_docx_path))
    theme_map = _load_theme_color_map(template_docx_path)
    for style in template_doc.styles:
        if getattr(style, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
            continue

        level = _extract_style_outline_level(style)
        if level is None or level < 1 or level > 6:
            continue

        rgb = _extract_rgb_from_style(style, theme_map)
        if rgb is None:
            continue
        if level not in color_map:
            color_map[level] = rgb

    return color_map


def _align_from_css(value: str) -> Optional[int]:
    v = value.strip().lower()
    if v == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if v == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if v == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if v in {"justify", "both"}:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return None


def _is_bold_value(value: str) -> Optional[bool]:
    v = value.strip().lower()
    if v in {"bold", "bolder"}:
        return True
    if v in {"normal", "lighter"}:
        return False
    if v.isdigit():
        return int(v) >= 600
    return None


def _is_italic_value(value: str) -> Optional[bool]:
    v = value.strip().lower()
    if v in {"italic", "oblique"}:
        return True
    if v == "normal":
        return False
    return None


def _is_underline_value(value: str) -> Optional[bool]:
    v = value.strip().lower()
    if "underline" in v:
        return True
    if v in {"none", "initial", "inherit"}:
        return False
    return None


def _style_font_size_pt(style_text: str) -> Optional[float]:
    match = re.search(r"font-size\s*:\s*([^;]+)", style_text or "", flags=re.IGNORECASE)
    if not match:
        return None
    value = match.group(1).strip().lower()
    pt_match = re.match(r"(-?\d+(?:\.\d+)?)pt$", value)
    if pt_match:
        return float(pt_match.group(1))
    px_match = re.match(r"(-?\d+(?:\.\d+)?)px$", value)
    if px_match:
        return float(px_match.group(1)) * 72.0 / 96.0
    return None


def _style_font_weight(style_text: str) -> Optional[bool]:
    match = re.search(r"font-weight\s*:\s*([^;]+)", style_text or "", flags=re.IGNORECASE)
    if not match:
        return None
    value = match.group(1).strip().lower()
    if value == "bold":
        return True
    if value.isdigit():
        return int(value) >= 600
    return None


def _style_text_align(style_text: str) -> Optional[str]:
    match = re.search(r"text-align\s*:\s*([^;]+)", style_text or "", flags=re.IGNORECASE)
    if not match:
        return None
    value = match.group(1).strip().lower()
    if value in {"left", "center", "right", "justify", "both"}:
        return "justify" if value == "both" else value
    return None


def _style_color_rgb(style_text: str) -> Optional[RGBColor]:
    match = re.search(r"color\s*:\s*([^;]+)", style_text or "", flags=re.IGNORECASE)
    if not match:
        return None
    value = match.group(1).strip().lower()
    hex_match = re.match(r"#([0-9a-f]{6})$", value)
    if hex_match:
        raw = hex_match.group(1)
        return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    rgb_match = re.match(r"rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)", value)
    if rgb_match:
        r = max(0, min(255, int(rgb_match.group(1))))
        g = max(0, min(255, int(rgb_match.group(2))))
        b = max(0, min(255, int(rgb_match.group(3))))
        return RGBColor(r, g, b)
    return None


def _parse_css_color_to_rgb(value: str) -> Optional[RGBColor]:
    raw = (value or "").strip().lower()
    hex_match = re.match(r"#([0-9a-f]{6})$", raw)
    if hex_match:
        token = hex_match.group(1)
        return RGBColor(int(token[0:2], 16), int(token[2:4], 16), int(token[4:6], 16))
    rgb_match = re.match(r"rgb\(\s*(\d{1,3})\s*,\s*(\d{1,3})\s*,\s*(\d{1,3})\s*\)$", raw)
    if rgb_match:
        r = max(0, min(255, int(rgb_match.group(1))))
        g = max(0, min(255, int(rgb_match.group(2))))
        b = max(0, min(255, int(rgb_match.group(3))))
        return RGBColor(r, g, b)
    return None


def _looks_like_list_item(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    return bool(
        re.match(r"^\d+\s*[、.．)]\s*", stripped)
        or re.match(r"^[（(]\d+[)）]\s*", stripped)
        or re.match(r"^[一二三四五六七八九十百千]+[、.．]\s*", stripped)
    )


def _looks_like_heading_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped or len(stripped) > 60:
        return False
    if any(ch in stripped for ch in ("。", "，", "；", "：", "?", "？", "!", "！")):
        return False
    return True


def _extract_heading_level_hint(text: str) -> Optional[int]:
    stripped = text.strip()
    if not stripped:
        return None
    if re.match(r"^第[一二三四五六七八九十百千0-9]+[章节篇部]\s*", stripped):
        return 1
    if re.match(r"^\d+\.\d+(?:\.\d+){0,2}\s*[、.．]?\s*", stripped):
        return 3
    if re.match(r"^[一二三四五六七八九十百千]+[、.．]\s*", stripped):
        return 2
    return None


def _heading_candidates_for_style(
    text: str,
    style_text: str,
    body_font_size_pt: float,
) -> Optional[int]:
    if not _looks_like_heading_text(text) or _looks_like_list_item(text):
        return None

    level_hint = _extract_heading_level_hint(text)
    if level_hint is not None:
        return level_hint

    font_size = _style_font_size_pt(style_text)
    if font_size is None:
        return None

    font_weight = _style_font_weight(style_text)
    align = _style_text_align(style_text)
    if len(text) <= 60 and not any(ch in text for ch in ("。", "，", "；", "：", "?", "？", "!", "！")):
        if align == "center" and font_size >= body_font_size_pt + 4:
            return 1
        if font_size >= body_font_size_pt + 5:
            return 1
        if font_size >= body_font_size_pt + 3:
            return 2
        if font_size >= body_font_size_pt + 1.2 and font_weight is True:
            return 3
        if font_size >= body_font_size_pt + 2 and font_weight is True and align == "center":
            return 2
    return None


def _add_runs_from_node(paragraph, node, style_state: Dict[str, bool]) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = style_state.get("bold", False)
        run.italic = style_state.get("italic", False)
        run.underline = style_state.get("underline", False)
        return

    if not isinstance(node, Tag):
        return

    local = dict(style_state)
    tag_name = node.name.lower()
    if tag_name in {"strong", "b"}:
        local["bold"] = True
    if tag_name in {"em", "i"}:
        local["italic"] = True
    if tag_name == "u":
        local["underline"] = True

    if tag_name == "br":
        paragraph.add_run("\n")
        return

    for child in node.children:
        _add_runs_from_node(paragraph, child, local)


def _apply_paragraph_style(paragraph, style: Dict[str, str], fallback: Dict[str, str]) -> None:
    merged = dict(fallback)
    merged.update(style)

    pfmt = paragraph.paragraph_format

    if "text-align" in merged:
        align = _align_from_css(merged["text-align"])
        if align is not None:
            paragraph.alignment = align

    if "line-height" in merged:
        lh = merged["line-height"].strip().lower()
        if re.match(r"^-?\d+(?:\.\d+)?$", lh):
            pfmt.line_spacing = float(lh)
        else:
            lh_pt = _parse_length_to_pt(lh)
            if lh_pt is not None:
                pfmt.line_spacing = Pt(lh_pt)

    if "margin-top" in merged:
        mt = _parse_length_to_pt(merged["margin-top"])
        if mt is not None:
            pfmt.space_before = Pt(mt)

    if "margin-bottom" in merged:
        mb = _parse_length_to_pt(merged["margin-bottom"])
        if mb is not None:
            pfmt.space_after = Pt(mb)

    if "text-indent" in merged:
        ti = _parse_length_to_pt(merged["text-indent"])
        if ti is not None:
            pfmt.first_line_indent = Pt(ti)

    font_name = merged.get("font-family")
    font_size_pt = _parse_length_to_pt(merged.get("font-size", "")) if "font-size" in merged else None
    bold = _is_bold_value(merged["font-weight"]) if "font-weight" in merged else None
    italic = _is_italic_value(merged["font-style"]) if "font-style" in merged else None
    underline = None
    if "text-decoration-line" in merged:
        underline = _is_underline_value(merged["text-decoration-line"])
    elif "text-decoration" in merged:
        underline = _is_underline_value(merged["text-decoration"])
    font_color = _parse_css_color_to_rgb(merged.get("color", "")) if "color" in merged else None

    for run in paragraph.runs:
        if font_name:
            first_family = font_name.split(",")[0].strip().strip('"\'')
            run.font.name = first_family
            run._element.rPr.rFonts.set(qn("w:eastAsia"), first_family)
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if font_color is not None:
            run.font.color.rgb = font_color


def _apply_heading_paragraph_style(
    doc: Document,
    paragraph,
    tag_name: str,
    heading_style_candidates: Dict[str, Tuple[str, ...]],
) -> bool:
    candidates = heading_style_candidates.get(tag_name, ())
    for style_name in candidates:
        try:
            paragraph.style = doc.styles[style_name]
            return True
        except Exception:
            continue

    # Fallback: keep the paragraph unstyled if the template does not expose a heading style.
    return False


def _force_outline_level(paragraph, level: int) -> None:
    p = paragraph._p
    ppr = p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(max(0, min(8, level - 1))))


def _set_paragraph_heading_level(
    doc: Document,
    paragraph,
    level: int,
    heading_style_candidates: Dict[str, Tuple[str, ...]],
) -> None:
    tag_name = f"h{max(1, min(6, level))}"
    _apply_heading_paragraph_style(doc, paragraph, tag_name, heading_style_candidates)
    _force_outline_level(paragraph, level)


def html_to_docx(
    html_path: Path,
    output_path: Path,
    template_docx_path: Optional[Path] = None,
) -> None:
    html = html_path.read_text(encoding="utf-8")
    page_layout, style_map = _extract_page_and_style(html)
    soup = BeautifulSoup(html, "html.parser")

    body = soup.body
    if body is None:
        raise ValueError("HTML missing <body> element")

    container = body.select_one(".page") or body

    doc = _create_output_document(template_docx_path)
    _apply_page_layout(doc, page_layout)
    _set_modern_compatibility_mode(doc, mode="16")
    heading_style_candidates = _build_heading_style_candidates_from_template(template_docx_path)
    heading_color_map = _extract_heading_color_map_from_template(template_docx_path)

    page_style = {}
    for selector in ("body.docx-layout-sync .page", "body .page", ".page"):
        if selector in style_map:
            page_style.update(style_map[selector])

    body_style = style_map.get("body.docx-layout-sync .page", {}) or style_map.get("body .page", {}) or style_map.get(".page", {})
    body_font_size_pt = _style_font_size_pt("; ".join(f"{k}: {v}" for k, v in body_style.items())) or 12.0

    for node in container.descendants:
        if not isinstance(node, Tag):
            continue
        name = node.name.lower()
        if name not in BLOCK_TAGS:
            continue
        # Skip nested block nodes, but keep direct children of the chosen container.
        if (
            node.parent
            and isinstance(node.parent, Tag)
            and node.parent is not container
            and node.parent.name.lower() in BLOCK_TAGS
        ):
            continue

        text_content = node.get_text("", strip=True)
        if not text_content:
            continue

        paragraph = doc.add_paragraph()
        paragraph_heading_level: Optional[int] = None
        if name in HEADING_TAG_LEVEL:
            paragraph_heading_level = HEADING_TAG_LEVEL[name]
            _set_paragraph_heading_level(
                doc,
                paragraph,
                paragraph_heading_level,
                heading_style_candidates,
            )
        _add_runs_from_node(paragraph, node, {"bold": False, "italic": False, "underline": False})
        tag_style = _combined_style(name, style_map)
        inline_style_text = node.get("style", "") or ""

        if name not in HEADING_TAG_LEVEL:
            merged_style_text = " ".join(
                [
                    inline_style_text,
                    " ".join(f"{k}: {v}" for k, v in tag_style.items()),
                ]
            ).strip()
            inferred_level = _heading_candidates_for_style(text_content, merged_style_text, body_font_size_pt)
            if inferred_level is not None:
                paragraph_heading_level = inferred_level
                _set_paragraph_heading_level(
                    doc,
                    paragraph,
                    inferred_level,
                    heading_style_candidates,
                )

        _apply_paragraph_style(paragraph, tag_style, page_style)

        # If HTML does not explicitly set color, enforce template heading color by level.
        if paragraph_heading_level is not None:
            explicit_color = _style_color_rgb(inline_style_text)
            css_color = _parse_css_color_to_rgb(tag_style.get("color", ""))
            if css_color is None:
                css_color = _parse_css_color_to_rgb(page_style.get("color", ""))
            if explicit_color is None and css_color is None:
                template_color = heading_color_map.get(paragraph_heading_level)
                if template_color is not None:
                    for run in paragraph.runs:
                        run.font.color.rgb = template_color

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    args = parse_args()
    if not args.html.exists():
        raise FileNotFoundError(f"HTML file not found: {args.html}")

    html_to_docx(args.html, args.output, template_docx_path=args.template_docx)
    print(f"DOCX saved to: {args.output}")


if __name__ == "__main__":
    main()
