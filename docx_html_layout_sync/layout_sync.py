"""Read DOCX layout settings and apply them as CSS to HTML content."""

from __future__ import annotations

import json
import re
import statistics
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Length


STYLE_ID = "docx-layout-sync-style"
HEADING_STYLE_PATTERNS = ("heading", "标题")
HEADING_TEXT_PATTERNS = (
    r"^第[一二三四五六七八九十百千0-9]+[章节篇部]\s*",
    r"^[一二三四五六七八九十百千]+[、.．]\s*",
    r"^[（(][一二三四五六七八九十百千0-9]+[)）]\s*",
    r"^\d+(?:\.\d+){0,3}[、.．]\s*",
)
TITLE_STYLE_CANDIDATES = ("Title", "标题", "文档标题")
BODY_STYLE_CANDIDATES = (
    "Normal",
    "正文",
    "Body Text",
    "正文文本",
    "Text Body",
)
TAG_STYLE_CANDIDATES = {
    "h1": ("Title", "标题", "文档标题", "Heading 1", "标题 1", "标题1"),
    "h2": ("Heading 1", "标题 1", "标题1", "Heading 2", "标题 2", "标题2"),
    "h3": ("Heading 2", "标题 2", "标题2", "Heading 3", "标题 3", "标题3"),
    "h4": ("Heading 3", "标题 3", "标题3", "Heading 4", "标题 4", "标题4"),
    "h5": ("Heading 4", "标题 4", "标题4", "Heading 5", "标题 5", "标题5"),
    "h6": ("Heading 5", "标题 5", "标题5", "Heading 6", "标题 6", "标题6"),
}
DEFAULT_FONT_STACK = "'宋体', 'SimSun', 'Microsoft YaHei', serif"
ACTIVE_THEME_COLOR_MAP: Dict[str, str] = {}


def _length_to_pt(value: Optional[Length]) -> Optional[float]:
    if value is None:
        return None
    return round(float(value.pt), 3)


def _format_pt(value: Optional[float], fallback: float) -> str:
    if value is None:
        value = fallback
    return f"{round(float(value), 3)}pt"


def _pt_to_mm(value: Optional[float]) -> Optional[float]:
    if value is None:
        return None
    return round(float(value) * 0.352778, 3)


def _css_font_family(font_name: Optional[str], fallback: str = "serif") -> str:
    if not font_name:
        return fallback
    escaped = font_name.replace("\\", "\\\\").replace("'", "\\'")
    return f"'{escaped}'"


def _alignment_to_css(value: Any) -> Optional[str]:
    if value is None:
        return None

    name = getattr(value, "name", None)
    if name:
        raw = name.strip().lower()
    else:
        raw = str(value).strip().lower()

    mapping = {
        "left": "left",
        "center": "center",
        "right": "right",
        "justify": "justify",
        "both": "justify",
        "distribute": "justify",
        "thai_distribute": "justify",
        "0": "left",
        "1": "center",
        "2": "right",
        "3": "justify",
    }
    return mapping.get(raw)


def _font_name_from_rfonts(rfonts: Any) -> Optional[str]:
    if rfonts is None:
        return None
    for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
        value = rfonts.get(qn(f"w:{attr}"))
        if value:
            return str(value)
    return None


def _font_size_pt_from_rpr(rpr: Any) -> Optional[float]:
    if rpr is None:
        return None
    sz = getattr(rpr, "sz", None)
    if sz is None:
        return None
    val = sz.val
    if val is None:
        return None
    try:
        return round(float(val) / 2.0, 3)
    except (TypeError, ValueError):
        return None


def _normalize_hex_color(value: str) -> Optional[str]:
    if not value:
        return None
    cleaned = value.strip().lstrip("#")
    if len(cleaned) != 6 or not re.fullmatch(r"[0-9a-fA-F]{6}", cleaned):
        return None
    return f"#{cleaned.upper()}"


def _normalize_theme_color_key(value: str) -> str:
    lowered = (value or "").strip().lower().replace("_", "")
    replacements = {
        "text1": "dk1",
        "background1": "lt1",
        "text2": "dk2",
        "background2": "lt2",
        "followedhyperlink": "folhlink",
    }
    return replacements.get(lowered, lowered)


def _load_theme_color_map(docx_path: Path) -> Dict[str, str]:
    result: Dict[str, str] = {}
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            theme_name = None
            for name in zf.namelist():
                if name.lower().startswith("word/theme/") and name.lower().endswith(".xml"):
                    theme_name = name
                    break
            if theme_name is None:
                return result

            xml_bytes = zf.read(theme_name)
            root = ET.fromstring(xml_bytes)
            ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            clr_scheme = root.find(".//a:clrScheme", ns)
            if clr_scheme is None:
                return result

            for child in list(clr_scheme):
                tag_name = child.tag.split("}")[-1]
                key = _normalize_theme_color_key(tag_name)
                srgb = child.find(".//a:srgbClr", ns)
                if srgb is not None:
                    hex_color = _normalize_hex_color(srgb.attrib.get("val", ""))
                    if hex_color:
                        result[key] = hex_color
                    continue

                sys_clr = child.find(".//a:sysClr", ns)
                if sys_clr is not None:
                    hex_color = _normalize_hex_color(sys_clr.attrib.get("lastClr", ""))
                    if hex_color:
                        result[key] = hex_color
    except Exception:
        return result

    return result


def _color_hex_from_color_obj(color_obj: Any) -> Optional[str]:
    if color_obj is None:
        return None
    rgb = getattr(color_obj, "rgb", None)
    if rgb is None:
        return None
    raw = str(rgb)
    return _normalize_hex_color(raw)


def _color_hex_from_rpr(rpr: Any) -> Optional[str]:
    if rpr is None:
        return None
    color = getattr(rpr, "color", None)
    if color is None:
        return None
    val = getattr(color, "val", None)
    if val is not None:
        normalized = _normalize_hex_color(str(val))
        if normalized:
            return normalized

    theme_color = color.get(qn("w:themeColor")) if hasattr(color, "get") else None
    if theme_color:
        key = _normalize_theme_color_key(str(theme_color))
        themed = ACTIVE_THEME_COLOR_MAP.get(key)
        if themed:
            return themed
    return None


def _resolve_style_font_name(style: Any, visited: Optional[Set[int]] = None) -> Optional[str]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    direct_name = getattr(getattr(style, "font", None), "name", None)
    if direct_name:
        return str(direct_name)

    rpr = getattr(getattr(style, "_element", None), "rPr", None)
    rfonts = getattr(rpr, "rFonts", None) if rpr is not None else None
    xml_name = _font_name_from_rfonts(rfonts)
    if xml_name:
        return xml_name

    return _resolve_style_font_name(getattr(style, "base_style", None), visited=visited)


def _resolve_style_font_size_pt(style: Any, visited: Optional[Set[int]] = None) -> Optional[float]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    direct_size = getattr(getattr(style, "font", None), "size", None)
    if direct_size is not None:
        return _length_to_pt(direct_size)

    rpr = getattr(getattr(style, "_element", None), "rPr", None)
    xml_size = _font_size_pt_from_rpr(rpr)
    if xml_size is not None:
        return xml_size

    return _resolve_style_font_size_pt(getattr(style, "base_style", None), visited=visited)


def _resolve_style_bold(style: Any, visited: Optional[Set[int]] = None) -> Optional[bool]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    direct_bold = getattr(getattr(style, "font", None), "bold", None)
    if direct_bold is not None:
        return bool(direct_bold)

    return _resolve_style_bold(getattr(style, "base_style", None), visited=visited)


def _resolve_style_italic(style: Any, visited: Optional[Set[int]] = None) -> Optional[bool]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    direct_italic = getattr(getattr(style, "font", None), "italic", None)
    if direct_italic is not None:
        return bool(direct_italic)

    return _resolve_style_italic(getattr(style, "base_style", None), visited=visited)


def _resolve_style_underline(style: Any, visited: Optional[Set[int]] = None) -> Optional[bool]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    direct_underline = getattr(getattr(style, "font", None), "underline", None)
    if direct_underline is not None:
        return bool(direct_underline)

    return _resolve_style_underline(getattr(style, "base_style", None), visited=visited)


def _resolve_style_color_hex(style: Any, visited: Optional[Set[int]] = None) -> Optional[str]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    direct_color = _color_hex_from_color_obj(getattr(getattr(style, "font", None), "color", None))
    if direct_color:
        return direct_color

    rpr = getattr(getattr(style, "_element", None), "rPr", None)
    xml_color = _color_hex_from_rpr(rpr)
    if xml_color:
        return xml_color

    return _resolve_style_color_hex(getattr(style, "base_style", None), visited=visited)


def _resolve_style_alignment(style: Any, visited: Optional[Set[int]] = None) -> Optional[str]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    paragraph_format = getattr(style, "paragraph_format", None)
    if paragraph_format is not None and paragraph_format.alignment is not None:
        css_align = _alignment_to_css(paragraph_format.alignment)
        if css_align:
            return css_align

    return _resolve_style_alignment(getattr(style, "base_style", None), visited=visited)


def _resolve_style_line_spacing(style: Any, visited: Optional[Set[int]] = None) -> Dict[str, Any]:
    if style is None:
        return {"line_spacing": None, "line_spacing_pt": None}
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return {"line_spacing": None, "line_spacing_pt": None}
    visited.add(style_id)

    paragraph_format = getattr(style, "paragraph_format", None)
    if paragraph_format is not None:
        line_spacing = paragraph_format.line_spacing
        if line_spacing is not None and hasattr(line_spacing, "pt"):
            return {
                "line_spacing": None,
                "line_spacing_pt": _length_to_pt(line_spacing),
            }
        if isinstance(line_spacing, (int, float)):
            return {"line_spacing": round(float(line_spacing), 3), "line_spacing_pt": None}

    return _resolve_style_line_spacing(getattr(style, "base_style", None), visited=visited)


def _resolve_style_space_before_after(style: Any, visited: Optional[Set[int]] = None) -> Dict[str, Any]:
    if style is None:
        return {"space_before_pt": None, "space_after_pt": None}
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return {"space_before_pt": None, "space_after_pt": None}
    visited.add(style_id)

    paragraph_format = getattr(style, "paragraph_format", None)
    before_pt = None
    after_pt = None
    if paragraph_format is not None:
        if paragraph_format.space_before is not None:
            before_pt = _length_to_pt(paragraph_format.space_before)
        if paragraph_format.space_after is not None:
            after_pt = _length_to_pt(paragraph_format.space_after)
        if before_pt is not None or after_pt is not None:
            return {"space_before_pt": before_pt, "space_after_pt": after_pt}

    return _resolve_style_space_before_after(
        getattr(style, "base_style", None),
        visited=visited,
    )


def _resolve_style_first_line_indent(style: Any, visited: Optional[Set[int]] = None) -> Optional[float]:
    if style is None:
        return None
    if visited is None:
        visited = set()
    style_id = id(style)
    if style_id in visited:
        return None
    visited.add(style_id)

    paragraph_format = getattr(style, "paragraph_format", None)
    if paragraph_format is not None and paragraph_format.first_line_indent is not None:
        return _length_to_pt(paragraph_format.first_line_indent)

    return _resolve_style_first_line_indent(getattr(style, "base_style", None), visited=visited)


def _find_style_by_name(doc: Document, style_name: str) -> Any:
    for style in doc.styles:
        if getattr(style, "name", "").lower() == style_name.lower():
            return style
    return None


def _extract_style_typography(style: Any, fallback: Dict[str, Any]) -> Dict[str, Any]:
    if style is None:
        return fallback.copy()

    extracted = fallback.copy()
    font_name = _resolve_style_font_name(style)
    if font_name:
        extracted["font_name"] = font_name

    font_size_pt = _resolve_style_font_size_pt(style)
    if font_size_pt is not None:
        extracted["font_size_pt"] = font_size_pt

    bold = _resolve_style_bold(style)
    if bold is not None:
        extracted["bold"] = bold

    italic = _resolve_style_italic(style)
    if italic is not None:
        extracted["italic"] = italic

    underline = _resolve_style_underline(style)
    if underline is not None:
        extracted["underline"] = underline

    color_hex = _resolve_style_color_hex(style)
    if color_hex:
        extracted["color_hex"] = color_hex

    alignment = _resolve_style_alignment(style)
    if alignment:
        extracted["alignment"] = alignment

    spacing_info = _resolve_style_line_spacing(style)
    if spacing_info["line_spacing"] is not None:
        extracted["line_spacing"] = spacing_info["line_spacing"]
    if spacing_info["line_spacing_pt"] is not None:
        extracted["line_spacing_pt"] = spacing_info["line_spacing_pt"]

    before_after_info = _resolve_style_space_before_after(style)
    if before_after_info["space_before_pt"] is not None:
        extracted["space_before_pt"] = before_after_info["space_before_pt"]
    if before_after_info["space_after_pt"] is not None:
        extracted["space_after_pt"] = before_after_info["space_after_pt"]

    first_line_indent_pt = _resolve_style_first_line_indent(style)
    if first_line_indent_pt is not None:
        extracted["first_line_indent_pt"] = first_line_indent_pt

    return extracted


def _is_heading_style_name(style_name: str) -> bool:
    lowered = (style_name or "").strip().lower()
    if not lowered:
        return False
    if "title" in lowered:
        return True
    return any(pattern in lowered for pattern in HEADING_STYLE_PATTERNS)


def _first_run_with_text(paragraph: Any) -> Any:
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _extract_run_font_name(run: Any, paragraph_style: Any) -> Optional[str]:
    if run is None:
        return _resolve_style_font_name(paragraph_style)

    run_font_name = getattr(getattr(run, "font", None), "name", None)
    if run_font_name:
        return str(run_font_name)

    rpr = getattr(getattr(run, "_element", None), "rPr", None)
    rfonts = getattr(rpr, "rFonts", None) if rpr is not None else None
    xml_font = _font_name_from_rfonts(rfonts)
    if xml_font:
        return xml_font

    return _resolve_style_font_name(paragraph_style)


def _extract_run_font_size_pt(run: Any, paragraph_style: Any) -> Optional[float]:
    if run is not None:
        run_size = getattr(getattr(run, "font", None), "size", None)
        if run_size is not None:
            return _length_to_pt(run_size)
        rpr = getattr(getattr(run, "_element", None), "rPr", None)
        xml_size = _font_size_pt_from_rpr(rpr)
        if xml_size is not None:
            return xml_size
    return _resolve_style_font_size_pt(paragraph_style)


def _extract_run_bold(
    run: Any,
    paragraph_style: Any,
    use_style_fallback: bool = True,
) -> Optional[bool]:
    if run is not None and run.bold is not None:
        return bool(run.bold)
    if use_style_fallback:
        return _resolve_style_bold(paragraph_style)
    return None


def _extract_run_italic(
    run: Any,
    paragraph_style: Any,
    use_style_fallback: bool = True,
) -> Optional[bool]:
    if run is not None and run.italic is not None:
        return bool(run.italic)
    if use_style_fallback:
        return _resolve_style_italic(paragraph_style)
    return None


def _extract_run_underline(
    run: Any,
    paragraph_style: Any,
    use_style_fallback: bool = True,
) -> Optional[bool]:
    if run is not None and run.underline is not None:
        return bool(run.underline)
    if use_style_fallback:
        return _resolve_style_underline(paragraph_style)
    return None


def _extract_run_color_hex(run: Any, paragraph_style: Any) -> Optional[str]:
    if run is not None:
        direct_color = _color_hex_from_color_obj(getattr(getattr(run, "font", None), "color", None))
        if direct_color:
            return direct_color

        rpr = getattr(getattr(run, "_element", None), "rPr", None)
        xml_color = _color_hex_from_rpr(rpr)
        if xml_color:
            return xml_color

    return _resolve_style_color_hex(paragraph_style)


def _extract_paragraph_alignment(paragraph: Any) -> Optional[str]:
    if paragraph.alignment is not None:
        return _alignment_to_css(paragraph.alignment)
    return _resolve_style_alignment(getattr(paragraph, "style", None))


def _extract_paragraph_line_spacing(paragraph: Any) -> Dict[str, Any]:
    paragraph_format = paragraph.paragraph_format
    line_spacing = paragraph_format.line_spacing
    if line_spacing is not None and hasattr(line_spacing, "pt"):
        return {"line_spacing": None, "line_spacing_pt": _length_to_pt(line_spacing)}
    if isinstance(line_spacing, (int, float)):
        return {"line_spacing": round(float(line_spacing), 3), "line_spacing_pt": None}
    return _resolve_style_line_spacing(getattr(paragraph, "style", None))


def _extract_paragraph_before_after(paragraph: Any) -> Dict[str, Any]:
    paragraph_format = paragraph.paragraph_format
    before_pt = _length_to_pt(paragraph_format.space_before)
    after_pt = _length_to_pt(paragraph_format.space_after)
    if before_pt is not None or after_pt is not None:
        return {"space_before_pt": before_pt, "space_after_pt": after_pt}
    return _resolve_style_space_before_after(getattr(paragraph, "style", None))


def _extract_paragraph_first_line_indent(paragraph: Any) -> Optional[float]:
    paragraph_format = paragraph.paragraph_format
    if paragraph_format.first_line_indent is not None:
        return _length_to_pt(paragraph_format.first_line_indent)
    return _resolve_style_first_line_indent(getattr(paragraph, "style", None))


def _looks_like_heading_text(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return any(re.match(pattern, stripped) for pattern in HEADING_TEXT_PATTERNS)


def _looks_like_body_sentence(text: str) -> bool:
    stripped = (text or "").strip()
    if len(stripped) >= 35:
        return True
    return any(p in stripped for p in ("。", "，", "；", "：", ",", ";", ":"))


def _extract_typography_from_paragraph(
    paragraph: Any,
    fallback: Dict[str, Any],
    use_style_bold_fallback: bool = True,
    use_style_italic_fallback: bool = True,
    use_style_underline_fallback: bool = True,
) -> Dict[str, Any]:
    extracted = fallback.copy()
    run = _first_run_with_text(paragraph)
    style = getattr(paragraph, "style", None)

    font_name = _extract_run_font_name(run, style)
    if font_name:
        extracted["font_name"] = font_name

    font_size_pt = _extract_run_font_size_pt(run, style)
    if font_size_pt is not None:
        extracted["font_size_pt"] = font_size_pt

    bold = _extract_run_bold(run, style, use_style_fallback=use_style_bold_fallback)
    if bold is not None:
        extracted["bold"] = bold
    elif not use_style_bold_fallback:
        extracted["bold"] = False

    italic = _extract_run_italic(run, style, use_style_fallback=use_style_italic_fallback)
    if italic is not None:
        extracted["italic"] = italic
    elif not use_style_italic_fallback:
        extracted["italic"] = False

    underline = _extract_run_underline(run, style, use_style_fallback=use_style_underline_fallback)
    if underline is not None:
        extracted["underline"] = underline
    elif not use_style_underline_fallback:
        extracted["underline"] = False

    color_hex = _extract_run_color_hex(run, style)
    if color_hex:
        extracted["color_hex"] = color_hex

    alignment = _extract_paragraph_alignment(paragraph)
    if alignment:
        extracted["alignment"] = alignment

    spacing_info = _extract_paragraph_line_spacing(paragraph)
    if spacing_info["line_spacing"] is not None:
        extracted["line_spacing"] = spacing_info["line_spacing"]
    if spacing_info["line_spacing_pt"] is not None:
        extracted["line_spacing_pt"] = spacing_info["line_spacing_pt"]

    before_after_info = _extract_paragraph_before_after(paragraph)
    if before_after_info["space_before_pt"] is not None:
        extracted["space_before_pt"] = before_after_info["space_before_pt"]
    if before_after_info["space_after_pt"] is not None:
        extracted["space_after_pt"] = before_after_info["space_after_pt"]

    first_line_indent_pt = _extract_paragraph_first_line_indent(paragraph)
    if first_line_indent_pt is not None:
        extracted["first_line_indent_pt"] = first_line_indent_pt

    return extracted


def _pick_body_paragraph(document: Document) -> Any:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        if _is_heading_style_name(style_name):
            continue
        if style_name and any(style_name.lower() == candidate.lower() for candidate in BODY_STYLE_CANDIDATES):
            return paragraph

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        if _is_heading_style_name(style_name):
            continue
        return paragraph
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            return paragraph
    return None


def _iter_body_paragraphs(document: Document) -> Any:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        if _is_heading_style_name(style_name):
            continue
        if _looks_like_heading_text(text):
            continue

        run = _first_run_with_text(paragraph)
        paragraph_size = _paragraph_size_pt(paragraph)
        is_bold = run is not None and run.bold is True
        # Headings often appear as short, bold, or oversized lines without explicit heading styles.
        if is_bold and len(text) <= 60 and not _looks_like_body_sentence(text):
            continue
        if paragraph_size is not None and paragraph_size >= 15.0 and len(text) <= 60:
            continue
        yield paragraph


def _iter_heading_paragraphs(document: Document, body_size_pt: float = 12.0) -> Any:
    size_threshold = max(14.0, body_size_pt + 0.8)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        if _is_heading_style_name(style_name):
            yield paragraph
            continue
        if _numbered_marker_kind(text) == "ordered-list":
            continue
        run = _first_run_with_text(paragraph)
        paragraph_size = _paragraph_size_pt(paragraph)
        if _looks_like_heading_text(text):
            yield paragraph
            continue
        if paragraph_size is not None and paragraph_size >= size_threshold and _is_short_heading_like_text(text):
            yield paragraph
            continue
        if run is not None and run.bold is True and _is_short_heading_like_text(text) and not _looks_like_body_sentence(text):
            yield paragraph
            continue


def _pick_title_paragraph(document: Document, body_typography: Dict[str, Any]) -> Any:
    body_size = float(body_typography.get("font_size_pt") or 12.0)
    size_threshold = max(16.0, body_size + 1.5)

    candidates = []
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        if style_name and "title" in style_name.lower():
            return paragraph

        text = paragraph.text.strip()
        paragraph_size = _paragraph_size_pt(paragraph) or 0.0
        run = _first_run_with_text(paragraph)
        alignment = _extract_paragraph_alignment(paragraph)

        score = 0
        if paragraph_size >= size_threshold:
            score += 4
        if _is_short_heading_like_text(text):
            score += 3
        if not _looks_like_body_sentence(text):
            score += 3
        if alignment == "center":
            score += 4
        if run is not None and run.bold is True:
            score += 1
        if _numbered_marker_kind(text) == "ordered-list":
            score -= 6
        if text.startswith("第"):
            score += 2
        candidates.append((score, paragraph))

    candidates.sort(key=lambda item: item[0], reverse=True)
    if candidates and candidates[0][0] >= 5:
        return candidates[0][1]
    return None


def _is_body_style_name(style_name: str) -> bool:
    lowered = (style_name or "").strip().lower()
    if not lowered:
        return False
    if _is_heading_style_name(lowered):
        return False
    return any(lowered == candidate.lower() for candidate in BODY_STYLE_CANDIDATES)


def _mode_or_default(values: Any, default: Any) -> Any:
    items = [v for v in values if v is not None]
    if not items:
        return default
    counts: Dict[Any, int] = {}
    for item in items:
        counts[item] = counts.get(item, 0) + 1
    # Keep deterministic behavior by preferring first seen item when counts tie.
    best = items[0]
    best_count = counts[best]
    for item in items:
        count = counts[item]
        if count > best_count:
            best = item
            best_count = count
    return best


def _median_or_default(values: Any, default: Optional[float]) -> Optional[float]:
    nums = [float(v) for v in values if isinstance(v, (int, float))]
    if not nums:
        return default
    return round(float(statistics.median(nums)), 3)


def _paragraph_size_pt(paragraph: Any) -> Optional[float]:
    run = _first_run_with_text(paragraph)
    if run is not None:
        if getattr(getattr(run, "font", None), "size", None) is not None:
            return _length_to_pt(getattr(run.font, "size"))
        rpr = getattr(getattr(run, "_element", None), "rPr", None)
        xml_size = _font_size_pt_from_rpr(rpr)
        if xml_size is not None:
            return xml_size
    return _resolve_style_font_size_pt(getattr(paragraph, "style", None))


def _refine_typography_from_paragraphs(
    document: Document,
    fallback: Dict[str, Any],
    paragraph_iter,
    bold_default: Optional[bool] = None,
    italic_default: Optional[bool] = None,
    underline_default: Optional[bool] = None,
) -> Dict[str, Any]:
    samples = []
    for paragraph in paragraph_iter(document):
        sample = _extract_typography_from_paragraph(
            paragraph,
            fallback,
            use_style_bold_fallback=False,
            use_style_italic_fallback=False,
            use_style_underline_fallback=False,
        )
        samples.append(sample)
        if len(samples) >= 120:
            break

    if not samples:
        refined = fallback.copy()
        if bold_default is not None:
            refined["bold"] = bold_default
        if italic_default is not None:
            refined["italic"] = italic_default
        if underline_default is not None:
            refined["underline"] = underline_default
        return refined

    refined = fallback.copy()
    refined["font_name"] = _mode_or_default([s.get("font_name") for s in samples], fallback.get("font_name"))
    refined["font_size_pt"] = _median_or_default([s.get("font_size_pt") for s in samples], fallback.get("font_size_pt"))
    refined["line_spacing"] = _median_or_default([s.get("line_spacing") for s in samples], fallback.get("line_spacing"))
    refined["line_spacing_pt"] = _median_or_default([s.get("line_spacing_pt") for s in samples], fallback.get("line_spacing_pt"))
    refined["space_before_pt"] = _median_or_default([s.get("space_before_pt") for s in samples], fallback.get("space_before_pt"))
    refined["space_after_pt"] = _median_or_default([s.get("space_after_pt") for s in samples], fallback.get("space_after_pt"))
    refined["alignment"] = _mode_or_default([s.get("alignment") for s in samples], fallback.get("alignment"))
    refined["bold"] = _mode_or_default([s.get("bold") for s in samples], bold_default if bold_default is not None else fallback.get("bold"))
    refined["italic"] = _mode_or_default(
        [s.get("italic") for s in samples],
        italic_default if italic_default is not None else fallback.get("italic"),
    )
    refined["underline"] = _mode_or_default(
        [s.get("underline") for s in samples],
        underline_default if underline_default is not None else fallback.get("underline"),
    )
    refined["color_hex"] = _mode_or_default([s.get("color_hex") for s in samples], fallback.get("color_hex"))
    refined["first_line_indent_pt"] = _median_or_default(
        [s.get("first_line_indent_pt") for s in samples],
        fallback.get("first_line_indent_pt"),
    )
    return refined


def _refine_body_typography_from_paragraphs(document: Document, fallback: Dict[str, Any]) -> Dict[str, Any]:
    return _refine_typography_from_paragraphs(
        document,
        fallback,
        _iter_body_paragraphs,
        bold_default=False,
        italic_default=False,
        underline_default=False,
    )


def _refine_heading_typography_from_paragraphs(document: Document, fallback: Dict[str, Any]) -> Dict[str, Any]:
    body_size = float(fallback.get("font_size_pt") or 12.0)
    return _refine_typography_from_paragraphs(
        document,
        fallback,
        lambda doc: _iter_heading_paragraphs(doc, body_size_pt=body_size),
        bold_default=None,
        italic_default=None,
        underline_default=None,
    )


def _doc_default_font_name(document: Document) -> Optional[str]:
    styles_element = getattr(document.styles, "element", None)
    if styles_element is None:
        return None
    doc_defaults = getattr(styles_element, "docDefaults", None)
    if doc_defaults is None:
        return None
    rpr_default = getattr(doc_defaults, "rPrDefault", None)
    if rpr_default is None:
        return None
    rpr = getattr(rpr_default, "rPr", None)
    if rpr is None:
        return None
    rfonts = getattr(rpr, "rFonts", None)
    return _font_name_from_rfonts(rfonts)


def _used_style_name_set(document: Document) -> Set[str]:
    names: Set[str] = set()
    for paragraph in document.paragraphs:
        style_name = getattr(getattr(paragraph, "style", None), "name", None)
        if style_name:
            names.add(style_name.strip().lower())
    return names


def _select_heading_style(document: Document, tag: str, used_style_names: Set[str]) -> Any:
    candidates = TAG_STYLE_CANDIDATES.get(tag, ())
    for name in candidates:
        style = _find_style_by_name(document, name)
        if style is not None and style.name.strip().lower() in used_style_names:
            return style
    return None


def extract_docx_layout_profile(docx_path: str) -> Dict[str, Any]:
    """Extract page and typography settings from DOCX."""
    docx_file = Path(docx_path)
    global ACTIVE_THEME_COLOR_MAP
    ACTIVE_THEME_COLOR_MAP = _load_theme_color_map(docx_file)
    document = Document(str(docx_file))
    section = document.sections[0]

    base_typography = {
        "font_name": None,
        "font_size_pt": 12.0,
        "line_spacing": 1.5,
        "line_spacing_pt": None,
        "space_before_pt": 0.0,
        "space_after_pt": 10.0,
        "bold": False,
        "italic": False,
        "underline": False,
        "alignment": "left",
        "color_hex": "#000000",
    }

    normal_style = _find_style_by_name(document, "Normal")
    body_typography = _extract_style_typography(normal_style, base_typography)
    body_typography = _refine_body_typography_from_paragraphs(document, body_typography)

    if not body_typography.get("font_name"):
        default_font_name = _doc_default_font_name(document)
        if default_font_name:
            body_typography["font_name"] = default_font_name

    used_style_names = _used_style_name_set(document)
    title_style = None
    for title_name in TITLE_STYLE_CANDIDATES:
        style = _find_style_by_name(document, title_name)
        if style is not None:
            title_style = style
            if style.name.strip().lower() in used_style_names:
                break

    headings: Dict[str, Dict[str, Any]] = {}
    heading_base = _refine_heading_typography_from_paragraphs(document, body_typography)

    for tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        heading_style = _select_heading_style(document, tag, used_style_names)
        headings[tag] = _extract_style_typography(heading_style, heading_base)

    title_paragraph = _pick_title_paragraph(document, body_typography)
    if title_paragraph is not None:
        headings["h1"] = _extract_typography_from_paragraph(title_paragraph, headings["h1"])
    elif title_style is not None:
        headings["h1"] = _extract_style_typography(title_style, headings["h1"])

    # Normalize only missing heading sizes, and never force bold.
    base_heading_size = headings["h2"].get("font_size_pt") or heading_base.get("font_size_pt") or body_typography.get("font_size_pt") or 16.0
    base_heading_size = float(base_heading_size)
    for index, tag in enumerate(("h2", "h3", "h4", "h5", "h6")):
        typography = headings[tag]
        if typography.get("font_size_pt") is None or typography.get("font_size_pt") == heading_base.get("font_size_pt"):
            typography["font_size_pt"] = max(11.0, round(base_heading_size - (index * 2.0), 3))

    profile = {
        "source_file": str(docx_file),
        "page": {
            "width_pt": _length_to_pt(section.page_width),
            "height_pt": _length_to_pt(section.page_height),
            "margin_top_pt": _length_to_pt(section.top_margin),
            "margin_right_pt": _length_to_pt(section.right_margin),
            "margin_bottom_pt": _length_to_pt(section.bottom_margin),
            "margin_left_pt": _length_to_pt(section.left_margin),
        },
        "body": body_typography,
        "headings": headings,
    }
    return profile


def _line_height_css(typography: Dict[str, Any]) -> str:
    line_spacing_pt = typography.get("line_spacing_pt")
    if line_spacing_pt is not None:
        return _format_pt(line_spacing_pt, fallback=18.0)

    line_spacing = typography.get("line_spacing")
    if line_spacing is None:
        return "1.5"
    return str(round(float(line_spacing), 3))


def build_css_from_profile(profile: Dict[str, Any]) -> str:
    """Build CSS string from extracted DOCX profile."""
    page = profile.get("page", {})
    body = profile.get("body", {})
    headings = profile.get("headings", {})

    width_pt = _format_pt(page.get("width_pt"), fallback=595.0)
    height_pt = _format_pt(page.get("height_pt"), fallback=842.0)
    margin_top_pt = _format_pt(page.get("margin_top_pt"), fallback=72.0)
    margin_right_pt = _format_pt(page.get("margin_right_pt"), fallback=72.0)
    margin_bottom_pt = _format_pt(page.get("margin_bottom_pt"), fallback=72.0)
    margin_left_pt = _format_pt(page.get("margin_left_pt"), fallback=72.0)

    body_font = _css_font_family(body.get("font_name"), fallback=DEFAULT_FONT_STACK)
    body_size_pt = _format_pt(body.get("font_size_pt"), fallback=12.0)
    body_line_height = _line_height_css(body)
    body_weight = "700" if body.get("bold") else "400"
    body_font_style = "italic" if body.get("italic") else "normal"
    body_text_decoration = "underline" if body.get("underline") else "none"
    body_color = body.get("color_hex") or "#000000"
    body_space_before = _format_pt(body.get("space_before_pt"), fallback=0.0)
    body_space_after = _format_pt(body.get("space_after_pt"), fallback=10.0)
    body_alignment = body.get("alignment") or "left"
    body_text_indent = _format_pt(body.get("first_line_indent_pt"), fallback=0.0)
    body_font_size_value = body.get("font_size_pt")
    if not isinstance(body_font_size_value, (int, float)):
        body_font_size_value = 12.0

    page_width_mm = _pt_to_mm(page.get("width_pt"))
    page_height_mm = _pt_to_mm(page.get("height_pt"))
    margin_top_mm = _pt_to_mm(page.get("margin_top_pt"))
    margin_right_mm = _pt_to_mm(page.get("margin_right_pt"))
    margin_bottom_mm = _pt_to_mm(page.get("margin_bottom_pt"))
    margin_left_mm = _pt_to_mm(page.get("margin_left_pt"))
    if page_width_mm is None:
        page_width_mm = 210.0
    if page_height_mm is None:
        page_height_mm = 297.0
    if margin_top_mm is None:
        margin_top_mm = 25.4
    if margin_right_mm is None:
        margin_right_mm = 25.4
    if margin_bottom_mm is None:
        margin_bottom_mm = 25.4
    if margin_left_mm is None:
        margin_left_mm = 25.4

    heading_rules = []
    for tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        typography = headings.get(tag, body)
        font_family = _css_font_family(typography.get("font_name"), fallback=body_font)
        font_size = _format_pt(
            typography.get("font_size_pt"),
            fallback=float(body_font_size_value),
        )
        line_height = _line_height_css(typography)
        font_weight = "700" if typography.get("bold") is True else "400"
        heading_font_style = "italic" if typography.get("italic") is True else "normal"
        heading_text_decoration = "underline" if typography.get("underline") is True else "none"
        heading_color = typography.get("color_hex") or body_color
        heading_space_before = _format_pt(
            typography.get("space_before_pt"),
            fallback=float(str(body.get("space_before_pt") or 0.0)),
        )
        heading_space_after = _format_pt(
            typography.get("space_after_pt"),
            fallback=float(str(body.get("space_after_pt") or 10.0)),
        )
        heading_alignment = typography.get("alignment") or body_alignment
        heading_text_indent = _format_pt(typography.get("first_line_indent_pt"), fallback=0.0)
        if tag == "h1":
            heading_text_indent = "0pt"
        heading_rules.append(
            "\n".join(
                [
                    f"body.docx-layout-sync {tag} {{",
                    f"  font-family: {font_family} !important;",
                    f"  font-size: {font_size} !important;",
                    f"  line-height: {line_height} !important;",
                    f"  font-weight: {font_weight} !important;",
                    f"  font-style: {heading_font_style} !important;",
                    f"  text-decoration-line: {heading_text_decoration} !important;",
                    f"  color: {heading_color} !important;",
                    f"  text-align: {heading_alignment} !important;",
                    f"  text-indent: {heading_text_indent} !important;",
                    f"  margin-top: {heading_space_before} !important;",
                    f"  margin-bottom: {heading_space_after} !important;",
                    "}",
                ]
            )
        )

    css_parts = [
        "/* Auto-generated from DOCX layout profile */",
        ":root {",
        f"  --docx-page-width: {width_pt};",
        f"  --docx-page-height: {height_pt};",
        f"  --docx-margin-top: {margin_top_pt};",
        f"  --docx-margin-right: {margin_right_pt};",
        f"  --docx-margin-bottom: {margin_bottom_pt};",
        f"  --docx-margin-left: {margin_left_pt};",
        f"  --docx-page-width-mm: {page_width_mm}mm;",
        f"  --docx-page-height-mm: {page_height_mm}mm;",
        f"  --docx-margin-top-mm: {margin_top_mm}mm;",
        f"  --docx-margin-right-mm: {margin_right_mm}mm;",
        f"  --docx-margin-bottom-mm: {margin_bottom_mm}mm;",
        f"  --docx-margin-left-mm: {margin_left_mm}mm;",
        "}",
        "@page {",
        "  size: var(--docx-page-width) var(--docx-page-height);",
        (
            "  margin: var(--docx-margin-top) var(--docx-margin-right) "
            "var(--docx-margin-bottom) var(--docx-margin-left);"
        ),
        "}",
        "html, body {",
        "  margin: 0;",
        "  padding: 0;",
        "}",
        "body.docx-layout-sync {",
        "  box-sizing: border-box;",
        "  width: auto !important;",
        "  min-height: 100vh !important;",
        "  margin: 0 !important;",
        "  padding: 20px 0 !important;",
        "  background: #f0f0f0 !important;",
        "}",
        "body.docx-layout-sync .page {",
        "  width: var(--docx-page-width-mm) !important;",
        "  min-height: var(--docx-page-height-mm) !important;",
        "  margin: 0 auto !important;",
        (
            "  padding: var(--docx-margin-top-mm) var(--docx-margin-right-mm) "
            "var(--docx-margin-bottom-mm) var(--docx-margin-left-mm) !important;"
        ),
        "  box-sizing: border-box;",
        f"  font-family: {body_font} !important;",
        f"  font-size: {body_size_pt} !important;",
        f"  line-height: {body_line_height} !important;",
        f"  font-weight: {body_weight} !important;",
        f"  font-style: {body_font_style} !important;",
        f"  text-decoration-line: {body_text_decoration} !important;",
        f"  text-align: {body_alignment} !important;",
        f"  color: {body_color} !important;",
        "  background: #ffffff !important;",
        "  box-shadow: 0 0 10px rgba(0, 0, 0, 0.15);",
        "  opacity: 1 !important;",
        "  visibility: visible !important;",
        "  overflow-wrap: break-word;",
        "}",
        "@media screen and (max-width: 900px) {",
        "  body.docx-layout-sync {",
        "    padding: 0 !important;",
        "  }",
        "  body.docx-layout-sync .page {",
        "    width: 100% !important;",
        "    min-height: auto !important;",
        "    margin: 0 !important;",
        "    box-shadow: none;",
        "  }",
        "}",
        "@media print {",
        "  body.docx-layout-sync {",
        "    background: none !important;",
        "    padding: 0 !important;",
        "  }",
        "  body.docx-layout-sync .page {",
        "    width: var(--docx-page-width-mm) !important;",
        "    min-height: var(--docx-page-height-mm) !important;",
        "    margin: 0 !important;",
        "    box-shadow: none;",
        "  }",
        "}",
        "body.docx-layout-sync p, body.docx-layout-sync li {",
        f"  margin-top: {body_space_before} !important;",
        f"  margin-bottom: {body_space_after} !important;",
        f"  text-indent: {body_text_indent} !important;",
        f"  font-style: {body_font_style} !important;",
        f"  text-decoration-line: {body_text_decoration} !important;",
        "}",
        "body.docx-layout-sync .page strong, body.docx-layout-sync .page b {",
        "  font-weight: inherit !important;",
        "}",
        "body.docx-layout-sync .page em, body.docx-layout-sync .page i {",
        "  font-style: inherit !important;",
        "}",
        "body.docx-layout-sync .page u {",
        "  text-decoration-line: inherit !important;",
        "  text-decoration: inherit !important;",
        "}",
    ]
    css_parts.extend(heading_rules)
    return "\n".join(css_parts) + "\n"


def _style_font_size_pt(style_text: str) -> Optional[float]:
    if not style_text:
        return None
    match = re.search(r"font-size\s*:\s*([^;]+)", style_text, flags=re.IGNORECASE)
    if not match:
        return None
    raw = match.group(1).strip().lower()
    pt_match = re.match(r"(-?\d+(?:\.\d+)?)pt$", raw)
    if pt_match:
        return float(pt_match.group(1))
    px_match = re.match(r"(-?\d+(?:\.\d+)?)px$", raw)
    if px_match:
        return float(px_match.group(1)) * 72.0 / 96.0
    return None


def _style_alignment(style_text: str) -> Optional[str]:
    lowered = (style_text or "").lower()
    if "text-align:center" in lowered or "text-align: center" in lowered:
        return "center"
    if "text-align:right" in lowered or "text-align: right" in lowered:
        return "right"
    if "text-align:justify" in lowered or "text-align: justify" in lowered:
        return "justify"
    if "text-align:left" in lowered or "text-align: left" in lowered:
        return "left"
    return None


def _text_heading_level(text: str) -> Optional[int]:
    stripped = text.strip()
    if not stripped:
        return None

    if re.match(r"^第[一二三四五六七八九十百千0-9]+[章节篇部]\s*", stripped):
        return 1
    if re.match(r"^\d+\.\d+(?:\.\d+){0,2}\s*[、.．]?\s*", stripped):
        # Multi-level numeric titles like 1.2 / 2.3.4 are commonly headings.
        return 3
    if re.match(r"^[一二三四五六七八九十百千]+[、.．]\s*", stripped):
        return 2
    return None


def _numbered_marker_kind(text: str) -> Optional[str]:
    stripped = text.strip()
    if not stripped:
        return None
    if re.match(r"^\d+\s*[、.．)]\s*", stripped):
        return "ordered-list"
    if re.match(r"^[（(]\d+[)）]\s*", stripped):
        return "ordered-list"
    if re.match(r"^[（(][一二三四五六七八九十百千]+[)）]\s*", stripped):
        return "ordered-list"
    if re.match(r"^\d+\.\d+(?:\.\d+){0,2}\s*[、.．]?\s*", stripped):
        return "section-number"
    return None


def _is_short_heading_like_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) > 48:
        return False
    return not any(ch in stripped for ch in ("。", "，", "；", "：", "?", "？", "!", "！"))


def _is_centered_like_title(style_text: str) -> bool:
    return _style_alignment(style_text) == "center"


def _is_body_sentence_like(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) >= 42:
        return True
    return any(ch in stripped for ch in ("。", "，", "；", "：", "?", "？", "!", "！"))


def _collect_paragraphs(root: Any) -> List[Tuple[Any, str]]:
    results: List[Tuple[Any, str]] = []
    for paragraph in root.find_all("p"):
        text = paragraph.get_text(" ", strip=True)
        if text:
            results.append((paragraph, text))
    return results


def _is_index_in_list_cluster(items: List[Tuple[Any, str]], index: int) -> bool:
    kind = _numbered_marker_kind(items[index][1])
    if kind != "ordered-list":
        return False

    neighbors = 0
    for step in (-1, 1):
        cursor = index + step
        while 0 <= cursor < len(items):
            other_kind = _numbered_marker_kind(items[cursor][1])
            if other_kind == "ordered-list":
                neighbors += 1
                cursor += step
                continue
            break
    return neighbors >= 1


def _promote_heading_like_paragraphs(html_text: str, profile: Dict[str, Any]) -> str:
    """Promote obvious heading-like <p> nodes into <h1>/<h2>/<h3> before CSS injection."""
    soup = BeautifulSoup(html_text, "html.parser")
    root = soup.body or soup

    body_size = float(profile.get("body", {}).get("font_size_pt") or 12.0)
    h1_profile = profile.get("headings", {}).get("h1", {})
    h2_profile = profile.get("headings", {}).get("h2", {})
    h3_profile = profile.get("headings", {}).get("h3", {})
    h1_size = float(h1_profile.get("font_size_pt") or max(body_size + 4.0, 16.0))
    h2_size = float(h2_profile.get("font_size_pt") or max(body_size + 2.0, 14.0))
    h3_size = float(h3_profile.get("font_size_pt") or max(body_size + 1.0, 13.0))

    paragraphs = _collect_paragraphs(root)
    promoted_h1 = False
    fallback_h1_index: Optional[int] = None
    fallback_h1_score = -10_000

    for index, (paragraph, text) in enumerate(paragraphs):

        explicit_level = _text_heading_level(text)
        marker_kind = _numbered_marker_kind(text)
        in_list_cluster = _is_index_in_list_cluster(paragraphs, index)
        inline_size = _style_font_size_pt(paragraph.get("style", ""))
        has_strong = paragraph.find(["strong", "b"]) is not None
        centered_like_title = _is_centered_like_title(paragraph.get("style", ""))
        short_like_heading = _is_short_heading_like_text(text)
        sentence_like = _is_body_sentence_like(text)

        level = explicit_level
        if level is not None and marker_kind == "ordered-list":
            # Plain ordered list markers such as 1、 2、 should remain body/list items.
            level = None
        if level is not None and in_list_cluster and marker_kind == "section-number":
            # Consecutive numbered lines are usually list items instead of section headings.
            level = None

        if level is None and marker_kind != "ordered-list" and short_like_heading and not sentence_like:
            if inline_size is not None and inline_size >= h1_size - 0.5:
                level = 1
            elif inline_size is not None and inline_size >= h2_size - 0.5:
                level = 2
            elif inline_size is not None and inline_size >= h3_size - 0.5:
                level = 3
            elif centered_like_title and index <= 4:
                level = 1 if index == 0 else 2
            elif has_strong and index <= 3:
                level = 2

        if level is None:
            if (
                marker_kind is None
                and index <= 5
                and short_like_heading
                and not sentence_like
            ):
                score = 0
                if index == 0:
                    score += 4
                if len(text) <= 28:
                    score += 3
                elif len(text) <= 48:
                    score += 1
                if centered_like_title:
                    score += 3
                if has_strong:
                    score += 1
                if inline_size is not None and inline_size >= h1_size - 0.5:
                    score += 4
                elif inline_size is not None and inline_size >= h2_size - 0.5:
                    score += 2
                if not sentence_like:
                    score += 2
                if score > fallback_h1_score:
                    fallback_h1_index = index
                    fallback_h1_score = score
            continue

        level = max(1, min(3, level))
        paragraph.name = f"h{level}"
        if level == 1:
            promoted_h1 = True

    if not promoted_h1 and fallback_h1_index is not None:
        candidate_paragraph, candidate_text = paragraphs[fallback_h1_index]
        if _numbered_marker_kind(candidate_text) is None:
            candidate_paragraph.name = "h1"

    return str(soup)


def inject_css_into_html(html_text: str, css_text: str) -> str:
    """Inject generated CSS into HTML and attach sync class to body."""
    style_tag = f'<style id="{STYLE_ID}">\n{css_text}</style>'
    has_body = re.search(r"<body\b[^>]*>", html_text, flags=re.IGNORECASE) is not None

    if not has_body:
        html_open = re.search(r"<html[^>]*>", html_text, flags=re.IGNORECASE)
        html_close = re.search(r"</html\s*>", html_text, flags=re.IGNORECASE)
        if html_open and html_close:
            head_close = re.search(r"</head\s*>", html_text, flags=re.IGNORECASE)
            if head_close and head_close.end() < html_close.start():
                body_start = head_close.end()
            else:
                body_start = html_open.end()
            html_text = (
                html_text[:body_start]
                + '\n<body class="docx-layout-sync">\n'
                + html_text[body_start:html_close.start()]
                + "\n</body>\n"
                + html_text[html_close.start():]
            )
        else:
            html_text = (
                "<!DOCTYPE html>\n<html>\n<head>\n</head>\n"
                '<body class="docx-layout-sync">\n'
                + html_text
                + "\n</body>\n</html>\n"
            )

    style_pattern = re.compile(
        rf"<style[^>]*id=[\"']{STYLE_ID}[\"'][^>]*>.*?</style>",
        flags=re.IGNORECASE | re.DOTALL,
    )
    if style_pattern.search(html_text):
        html_text = style_pattern.sub(style_tag, html_text)
    else:
        head_close = re.search(r"</head\s*>", html_text, flags=re.IGNORECASE)
        if head_close:
            insert_at = head_close.start()
            html_text = html_text[:insert_at] + style_tag + "\n" + html_text[insert_at:]
        else:
            html_open = re.search(r"<html[^>]*>", html_text, flags=re.IGNORECASE)
            if html_open:
                insert_at = html_open.end()
                html_text = (
                    html_text[:insert_at]
                    + "\n<head>\n"
                    + style_tag
                    + "\n</head>\n"
                    + html_text[insert_at:]
                )
            else:
                html_text = (
                    "<!DOCTYPE html>\n<html>\n<head>\n"
                    + style_tag
                    + "\n</head>\n<body>\n"
                    + html_text
                    + "\n</body>\n</html>\n"
                )

    body_open = re.search(r"<body\b([^>]*)>", html_text, flags=re.IGNORECASE)
    if body_open:
        attrs = body_open.group(1)
        class_match = re.search(r'class\s*=\s*["\']([^"\']*)["\']', attrs, flags=re.IGNORECASE)
        if class_match:
            classes = class_match.group(1).split()
            if "docx-layout-sync" not in classes:
                classes.append("docx-layout-sync")
                new_attr = f'class="{" ".join(classes)}"'
                attrs = (
                    attrs[: class_match.start()] + new_attr + attrs[class_match.end() :]
                )
        else:
            attrs += ' class="docx-layout-sync"'
        html_text = (
            html_text[: body_open.start()]
            + f"<body{attrs}>"
            + html_text[body_open.end() :]
        )

    body_open = re.search(r"<body\b[^>]*>", html_text, flags=re.IGNORECASE)
    body_close = re.search(r"</body\s*>", html_text, flags=re.IGNORECASE)
    if body_open and body_close:
        inner_html = html_text[body_open.end() : body_close.start()]
        if 'class="page"' not in inner_html and "class='page'" not in inner_html:
            html_text = (
                html_text[: body_open.end()]
                + '\n<div class="page">\n'
                + inner_html.strip()
                + '\n</div>\n'
                + html_text[body_close.start() :]
            )
    return html_text


def sync_docx_layout_to_html(
    docx_path: str,
    html_path: str,
    output_html_path: str,
    profile_output_path: Optional[str] = None,
) -> Dict[str, str]:
    """Apply DOCX layout settings to an HTML file and save output."""
    profile = extract_docx_layout_profile(docx_path)
    css = build_css_from_profile(profile)

    html_text = Path(html_path).read_text(encoding="utf-8")
    html_text = _promote_heading_like_paragraphs(html_text, profile)
    styled_html = inject_css_into_html(html_text, css)

    output_path = Path(output_html_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(styled_html, encoding="utf-8")

    profile_path = None
    if profile_output_path:
        profile_path = Path(profile_output_path)
        profile_path.parent.mkdir(parents=True, exist_ok=True)
        profile_path.write_text(
            json.dumps(profile, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    result = {
        "output_html": str(output_path.resolve()),
        "profile_json": str(profile_path.resolve()) if profile_path else "",
    }
    return result
